from pydoc import doc

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_resume():
    doc = Document()
    
    # --- PAGE & STYLE SETUP ---
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.4)
    section.left_margin = section.right_margin = Inches(0.25)

    # --- ADD FOOTER WITH PAGE NUMBERING ---
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.text = "Page "
    
    # Add current page number field
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    run = footer_p.add_run()
    fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
    instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w')))
    fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    
    footer_p.add_run(" of ")
    
    # Add total page count field
    run = footer_p.add_run()
    fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
    instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> NUMPAGES </w:instrText>'.format(nsdecls('w')))
    fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = pf.space_after = Pt(0)

    # Helper: Add hyperlink to a paragraph
    def add_hyperlink(paragraph, url, text):
        # Create a new relationship for the hyperlink
        rel_key = paragraph.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # Create hyperlink run
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor(0, 0, 255)  # blue
        run.underline = True
        
        # Add hyperlink XML element
        rPr = run._element.get_or_add_rPr()
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        
        # Add the actual hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), rel_key)
        run._element.addprevious(hyperlink)
        hyperlink.append(run._element)

    # Helper: Full-row Underline with Padding
    # `color` may be a hex string ('FF0000'), named color, or 'auto'.
    # `thickness` is the line width in eighths of a point (w:sz value).
    # Typical values: 4 (0.5pt), 8 (1pt), 16 (2pt). Increase for a bolder line.
    def add_section_underline(paragraph, color='4169e1', thickness=24):
        pBr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(thickness))
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), color)    
        pBdr.append(bottom)
        pBr.append(pBdr)
        #shd = OxmlElement('w:shd')
        #shd.set(qn('w:val'), 'clear')
        #shd.set(qn('w:color'), color)
        #shd.set(qn('w:fill'), color)     # the background colour
        #pBr.append(shd)

    
    def load_pii(pii_path=None):
        """Load personal info from a JSON file and return (name, email, phone, location).

        If `pii_path` is None, the function looks for `pii.json` next to this file.
        On failure (file missing or invalid JSON), returns empty strings.
        """
        import os, json

        if pii_path is None:
            pii_path = os.path.join(os.path.dirname(__file__), 'pii.json')

        try:
            with open(pii_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except Exception:
            return '', '', '', ''

        name = data.get('name', '')
        email = data.get('email', '')
        phone = data.get('phone', '')
        location = data.get('location', '')
        return name, email, phone, location

    # Load personal info from pii.json (falls back to empty strings on error)
    name, email, phone, location = load_pii()

    # --- CONDENSED CENTERED HEADER --- 
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_p.paragraph_format.space_after = Pt(1)
    name_run = name_p.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(28)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contact_p.paragraph_format.space_after = Pt(2)
    contact_run = contact_p.add_run(f"✉ {email}  | ✆ {phone}  | 📍 {location}")
    contact_run.font.size = Pt(12)

    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    link_p.add_run("</> ")
    add_hyperlink(link_p, "https://github.com/Michaelis105", "github.com/Michaelis105")
    link_p.add_run("  | [in] ")
    add_hyperlink(link_p, "https://linkedin.com/in/louiemichael", "linkedin.com/in/louiemichael")

    def add_page_break():
        page_break = doc.add_paragraph("")
        page_break.add_run().add_break(WD_BREAK.PAGE)

    def add_section_heading(text):
        sec = doc.add_paragraph(text)
        sec.runs[0].bold = True
        sec.runs[0].font.size = Pt(12)
        sec.paragraph_format.space_before = Pt(14)
        sec.paragraph_format.space_after = Pt(12)
        add_section_underline(sec)
 
    # --- SECTION: SUMMARY ---
    add_section_heading("SUMMARY")

    summary_text = "Full-stack software engineer with nearly 10 years of developing resilient, distributed AWS cloud systems in heavily-regulated environments at scale. Proven track record of steering engineering teams in building patented self-service banking technology supporting over $20 million in transactions and intelligent fleet management service overseeing 1000s of customer-facing financial devices. Combining deep full-stack expertise with the latest generative AI engineering technologies."
    summary_p1 = doc.add_paragraph(summary_text)
    summary_p1.paragraph_format.space_before = Pt(10)
    summary_p1.paragraph_format.line_spacing = 1.15

     # --- SECTION: TECHNICAL EXPERTISE ---
    add_section_heading("TECHNICAL EXPERTISE")

    skills = [
        ("Languages & Core Tech", "Python + Flask, Node.js + Express, Bash, Java + Spring, Typescript"),
        ("Frontend", "React, Vue, HTML, CSS, JavaScript"),
        ("Cloud & Infrastructure", "Amazon Web Services (AWS), Unix, Linux, Docker, Terraform, Git, GitHub"),
        ("Data Engineering", "SQL, NoSQL, MySQL, PostgreSQL, Snowflake, Kafka"),
        ("Observability/SRE", "Splunk, New Relic, Cloudwatch, PagerDuty, Playbooks, Technical Documentation"),
        ("Testing", "Jest, Cypress, Cucumber"),
        ("Architecture/System Design", "RESTful API, Distributed Systems, Event-Driven, Real-time, Microservices"),
        ("Leadership/Team Management", "Agile, Scrum, Cross-functional Team Coordination, Presentations All Audiences"),
        ("Growth", "Patent Process, Root-cause-analysis, Post-Mortems, Mentoring, Hackathons, Interviewing"),
        ("Generative AI Engineering", "Ollama, LangChain, ChromaDB, Windsurf, Copilot, Gemini"),
        ("Others", "3D Printing, Computer Aided Design (CAD), OnShape"),
    ]
    for skill in skills:
        p = doc.add_paragraph()
        category, items = skill
        run = p.add_run(f"• {category}")
        run.bold = True
        p.add_run(f": {items}")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    # --- SECTION: PROFESSIONAL WORK EXPERIENCE ---
    add_section_heading("PROFESSIONAL WORK EXPERIENCE")

    # helper to format company line with dates
    def add_company_heading(company, dates=None):
        company_p = doc.add_paragraph()
        company_p.paragraph_format.space_before = Pt(13)
        company_p.paragraph_format.tab_stops.add_tab_stop(Inches(8.0), WD_TAB_ALIGNMENT.RIGHT)
        run = company_p.add_run(company)
        run.bold = True
        run.font.size = Pt(11)
        if dates is not None:
            company_p.add_run(f"\t{dates}")
        add_section_underline(company_p, color='000000', thickness=1)

    def add_job(title, dates, bullets, level2_range=None):        
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(16)
        p2.paragraph_format.space_after = Pt(12)
        p2.paragraph_format.tab_stops.add_tab_stop(Inches(8.0), WD_TAB_ALIGNMENT.RIGHT)
        # separate runs so only job title is italic
        title_run = p2.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(11)
        if dates:
            date_run = p2.add_run(f"\t{dates}")
            date_run.font.size = Pt(11)
        
        for idx, b in enumerate(bullets):
            bp = doc.add_paragraph()
            # Handle both plain strings and tuples with formatting
            if isinstance(b, tuple):
                bold_text, is_bold, rest = b
                run = bp.add_run(f"• ")
                run = bp.add_run(f"{bold_text}")
                run.bold = is_bold
                bp.add_run(f" {rest}")
            else:
                bp.add_run(f"• {b}")

            bp.paragraph_format.space_before = Pt(4)
            bp.paragraph_format.space_after = Pt(4)

            # Indenting bullets slightly from the left margin
            if level2_range and level2_range[0] <= idx <= level2_range[1]:
                bp.paragraph_format.left_indent = Inches(0.6) # Deeper nested indent
                bp.paragraph_format.first_line_indent = Inches(-0.2)
            else:
                bp.paragraph_format.left_indent = Inches(0.4) # Standard slight indent
                bp.paragraph_format.first_line_indent = Inches(-0.2)

    # Capital One 
    c1_lead_bullets = [
        "Tech Lead launching U.S.-first patented self-service cashier’s check kiosk through nationwide roll-out.",
        "Supporting >$20 million in high-stakes transactions via customer cashier’s checks across money markets.",
        "Pruning 30% duplicate work via standardized MERN serverless tech stack across multiple self-service platforms.",
        "Steering team of seven engineers via pair programming, code reviews, and resiliency on-call playbooks.",
        "Conveying technical intent with product, engineers, and designers via architecture/dataflow/API design diagrams.",
        "Tech Lead for owning technical direction of green-field self-service instant payment issuance card kiosk.",
        "Engineering a LangChain and ChromaDB RAG-based prototype streamlining bank policy/procedure research.",
    ]

    # Capital One 
    c1_pa_bullets = [
        "Minimizing manual cloud deployments and version drift via Terraform-like infra-as-code managed CICD changes.",
        "Centralizing transaction and kiosk state at single source data lake via Kafka for real-time monitoring.",
        "Reducing kiosk deployment times by 80% by pioneering fleet management pub-sub operation code mechanism.",
    ]

    # Capital One 
    c1_sa_bullets = [
        "Developing ATM fleet managing/monitoring distributed system serving real-time operations/auditing.",
        "Building MSI to streamline ATM software platform lifecycle management, reducing per kiosk downtime by 60%.",
    ]

    add_company_heading("Capital One Financial")
    add_job("Lead Software Engineer – Bank Tech, Consumer Self-Servicing", "August 2024 – Present", c1_lead_bullets)
    add_job("Senior Software Engineer – Bank Tech, Associate In-Person Experience", "July 2021 – August 2024", c1_pa_bullets)
    add_job("Software Engineer – Retail Bank Tech, Digital Customer Experience", "July 2019 – July 2021", c1_sa_bullets)

    bloomberg_bullets = ["Decreased customer secure access outages by 10% via preemptive SAML certificate expiration notifications.",]
    add_company_heading("Bloomberg Industry Group", "August 2018 – July 2019")
    add_job("Software Engineer – Subscription Management and Customer Support Platform", None, bloomberg_bullets)
    
    vs_bullets = [
        "Optimized DotGov domain management web portal with customer service and GSA user feedback.",
        "Developed internal code dependency analysis reporting tool to analyze and report code security vulnerabilities."
    ]
    add_company_heading("Verisign, Inc.", "February 2017 – August 2018")
    add_job("Software Engineer I-II – Consolidated Top-Level Domain, Infrastructure Services", None, vs_bullets)
    
    lm_bullets = ["Modernized submarine sonar applications stack via integrating scaling and management Mesos/Marathon COTS.", "Introduced Docker containerization for hosting submarine applications."]
    add_company_heading("Lockheed Martin", "June 2016 – February 2017")
    add_job("Software Engineer Associate – Acoustic Rapid COTS Insertion System Services", None, lm_bullets)

    # --- SECTION: PATENTS --- 
    add_section_heading("PATENTS")
    
    patents = ["Systems and Methods for Securely Generating and Printing a Document (US20220414641A1)", "Graphical User Interface for Centralized Register Device Management and Monitoring (Notice of Allowance)"]
    for patent in patents:
        p = doc.add_paragraph(f"• {patent}")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    # --- SECTION: CERTIFICATIONS --- 
    add_section_heading("CERTIFICATIONS")
    
    certs = ["AWS Certified Developer Associate and Cloud Practitioner", "AWS Certified Generative AI Developer – Professional and Solutions Architect (scheduled Q2 2026)", "CompTIA Network+ Certification N10-006"]
    for cert in certs:
        p = doc.add_paragraph(f"• {cert}")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    # --- SECTION: EDUCATION --- 
    add_section_heading("EDUCATION")

    # Georgia Tech 
    pe1 = doc.add_paragraph()
    pe1.paragraph_format.space_before = Pt(8)
    pe1.paragraph_format.tab_stops.add_tab_stop(Inches(8.0), WD_TAB_ALIGNMENT.RIGHT)
    pe1.add_run("Georgia Institute of Technology").bold = True
    pe1.add_run("\tSpring 2021")
    pe1_sub = doc.add_paragraph()
    pe1_sub.paragraph_format.tab_stops.add_tab_stop(Inches(8.0), WD_TAB_ALIGNMENT.RIGHT)
    pe1_sub.add_run("M.S. in Computer Science (Computing Systems)").italic = True
    pe1_sub.add_run("\tAtlanta, GA")
    pe1_sub.paragraph_format.space_after = Pt(4)

    # Virginia Tech 
    pe2 = doc.add_paragraph()
    pe2.paragraph_format.space_before = Pt(8)
    pe2.paragraph_format.tab_stops.add_tab_stop(Inches(8.0), WD_TAB_ALIGNMENT.RIGHT)
    pe2.add_run("Virginia Polytechnic Institute and State University").bold = True
    pe2.add_run("\tMay 2016")
    pe2_sub = doc.add_paragraph()
    pe2_sub.paragraph_format.tab_stops.add_tab_stop(Inches(8.0), WD_TAB_ALIGNMENT.RIGHT)
    pe2_sub.add_run("B.S. in Computer Science").italic = True
    pe2_sub.add_run("\tBlacksburg, VA")
    pe2_sub.paragraph_format.space_after = Pt(4)

    doc.save('michael-louie-resume.docx')

create_resume()